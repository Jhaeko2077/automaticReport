from __future__ import annotations

from dataclasses import dataclass
import re
from pathlib import Path

from docx import Document

PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
QUESTION_RE = re.compile(r"pregunta\s*\d+\s*[:：]", re.IGNORECASE)


@dataclass
class QuestionSlot:
    question: str
    table_idx: int
    question_row_idx: int
    answer_row_idx: int
    col_idx: int


@dataclass
class DocAnalysis:
    placeholders: dict[str, str]
    questions: list[QuestionSlot]
    has_summary_section: bool
    has_diagram_section: bool


def _normalize(text: str) -> str:
    return " ".join(text.split()).strip().lower()


def _label_matches_cell(label: str, normalized_text: str, key: str) -> bool:
    if key == "student_id":
        compact = normalized_text.replace(" ", "")
        return compact in {"id", "id:", "id."}
    return label in normalized_text


def _replace_student_line(raw_text: str, value: str, key: str) -> str:
    lines = raw_text.splitlines()
    if not lines:
        return raw_text

    for idx, line in enumerate(lines):
        normalized = _normalize(line)
        for label, mapped_key in {
            "apellidos y nombres": "student_name",
            "id": "student_id",
            "dirección zonal/cfp": "student_address",
            "direccion zonal/cfp": "student_address",
            "carrera": "student_career",
            "curso/ mod. formativo": "student_course",
            "curso/mod. formativo": "student_course",
            "tema de trabajo final": "student_topic",
        }.items():
            if mapped_key != key:
                continue
            if not _label_matches_cell(label, normalized, key):
                continue

            prefix = line.split(":", 1)[0].strip() if ":" in line else line.strip()
            lines[idx] = f"{prefix}: {value}"
            return "\n".join(lines)

    return raw_text


def _strip_question_label(text: str) -> str:
    return QUESTION_RE.sub("", text).strip()


def _is_question_text(text: str) -> bool:
    normalized = _normalize(text)
    return "pregunta" in normalized and "?" in normalized


def _extract_question_from_row(row) -> tuple[str | None, int]:
    row_cells = [(cell.text or "").strip() for cell in row.cells]
    row_text = " ".join(part for part in row_cells if part)
    if not row_text:
        return None, 0

    match = QUESTION_RE.search(row_text)
    if not match:
        return None, 0

    question = row_text[match.end() :].strip()
    if not question:
        return None, 0

    # Intentamos ubicar la columna más probable para la respuesta.
    # Prioridad: celda con signo de pregunta, luego celda con texto más largo.
    col_idx = 0
    cells_with_q = [idx for idx, cell_text in enumerate(row_cells) if "?" in cell_text]
    if cells_with_q:
        col_idx = cells_with_q[0]
    else:
        non_empty_cells = [(idx, len(text)) for idx, text in enumerate(row_cells) if text]
        if non_empty_cells:
            col_idx = max(non_empty_cells, key=lambda item: item[1])[0]

    return question, col_idx


def analyze_docx(docx_path: str) -> DocAnalysis:
    doc = Document(docx_path)
    placeholders: dict[str, str] = {}
    questions: list[QuestionSlot] = []
    has_summary_section = False
    has_diagram_section = False

    for p in doc.paragraphs:
        text = p.text or ""
        for match in PLACEHOLDER_RE.finditer(text):
            key = match.group(1)
            placeholders.setdefault(key, text.strip())

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            row_question, row_col_idx = _extract_question_from_row(row)
            if row_question and row_idx + 1 < len(table.rows):
                answer_cell = table.rows[row_idx + 1].cells[row_col_idx]
                answer_text = (answer_cell.text or "").strip()
                next_row_question, _ = _extract_question_from_row(table.rows[row_idx + 1])
                if not _is_question_text(answer_text) and not next_row_question:
                    questions.append(
                        QuestionSlot(
                            question=row_question,
                            table_idx=table_idx,
                            question_row_idx=row_idx,
                            answer_row_idx=row_idx + 1,
                            col_idx=row_col_idx,
                        )
                    )

            for col_idx, cell in enumerate(row.cells):
                text = (cell.text or "").strip()
                for match in PLACEHOLDER_RE.finditer(text):
                    key = match.group(1)
                    placeholders.setdefault(key, text)

                normalized = _normalize(text)
                if normalized.startswith("resumen"):
                    has_summary_section = True
                if normalized.startswith("diagrama"):
                    has_diagram_section = True

                # Compatibilidad adicional: pregunta contenida íntegramente en una sola celda.
                if _is_question_text(text) and row_idx + 1 < len(table.rows):
                    answer_cell = table.rows[row_idx + 1].cells[col_idx]
                    answer_text = (answer_cell.text or "").strip()
                    if not _is_question_text(answer_text):
                        question_text = _strip_question_label(text)
                        if not any(
                            q.table_idx == table_idx and q.question_row_idx == row_idx and q.col_idx == col_idx
                            for q in questions
                        ):
                            questions.append(
                                QuestionSlot(
                                    question=question_text,
                                    table_idx=table_idx,
                                    question_row_idx=row_idx,
                                    answer_row_idx=row_idx + 1,
                                    col_idx=col_idx,
                                )
                            )

    return DocAnalysis(
        placeholders=placeholders,
        questions=questions,
        has_summary_section=has_summary_section,
        has_diagram_section=has_diagram_section,
    )


def extract_placeholders(docx_path: str) -> dict[str, str]:
    return analyze_docx(docx_path).placeholders


def fill_docx_template(docx_input: str, docx_output: str, replacements: dict[str, str]) -> None:
    doc = Document(docx_input)

    def replace_text(text: str) -> str:
        for key, value in replacements.items():
            text = text.replace(f"{{{{{key}}}}}", value)
        return text

    for p in doc.paragraphs:
        if p.text:
            p.text = replace_text(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    cell.text = replace_text(cell.text)

    output_path = Path(docx_output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def _split_non_empty_lines(value: str) -> list[str]:
    return [line.strip(" -\t") for line in value.splitlines() if line.strip(" -\t")]


def _find_first_empty_row(table, start_idx: int = 0) -> int | None:
    for row_idx in range(max(0, start_idx), len(table.rows)):
        row = table.rows[row_idx]
        if all(not (cell.text or "").strip() for cell in row.cells):
            return row_idx
    return None


def _table_contains_label(table, label: str) -> bool:
    normalized_label = _normalize(label)
    for row in table.rows:
        for cell in row.cells:
            if normalized_label in _normalize(cell.text or ""):
                return True
    return False


def _build_resource_items(raw_value: str) -> list[tuple[str, str]]:
    items: list[tuple[str, str]] = []
    for line in _split_non_empty_lines(raw_value):
        if "|" in line:
            left, right = line.split("|", 1)
            items.append((left.strip(), right.strip()))
            continue

        qty_match = re.match(r"^(.*?)(?:\s*[:\-]\s*|\s*\()\s*(\d+)\)?\s*$", line)
        if qty_match:
            items.append((qty_match.group(1).strip(), qty_match.group(2).strip()))
        else:
            items.append((line.strip(), "1"))
    return items


def _fill_resource_table(table, raw_value: str) -> bool:
    items = _build_resource_items(raw_value)
    if not items:
        return False

    start_row = _find_first_empty_row(table, start_idx=1)
    if start_row is None:
        return False

    max_rows = len(table.rows) - start_row
    for offset, (description, quantity) in enumerate(items[:max_rows]):
        row = table.rows[start_row + offset]
        if len(row.cells) >= 1:
            row.cells[0].text = description
        if len(row.cells) >= 2:
            row.cells[1].text = quantity
    return True


def _fill_schedule_table(table, raw_value: str) -> bool:
    activities = _split_non_empty_lines(raw_value)
    if not activities:
        return False

    start_row = _find_first_empty_row(table, start_idx=2)
    if start_row is None:
        return False

    for idx, activity in enumerate(activities):
        row_idx = start_row + idx
        if row_idx >= len(table.rows):
            break
        row = table.rows[row_idx]
        col_count = len(row.cells)
        if col_count < 2:
            continue

        row.cells[0].text = str(idx + 1)
        row.cells[1].text = activity

        # Marca de cronograma rotativa sobre las columnas disponibles.
        if col_count > 2:
            mark_col = 2 + (idx % (col_count - 2))
            row.cells[mark_col].text = "X"
    return True


def _fill_execution_table(table, operations: str, standards: str) -> bool:
    op_lines = _split_non_empty_lines(operations)
    std_lines = _split_non_empty_lines(standards)
    if not op_lines and not std_lines:
        return False

    start_row = _find_first_empty_row(table, start_idx=1)
    if start_row is None:
        return False

    target_rows = len(table.rows) - start_row
    total_lines = min(target_rows, max(len(op_lines), len(std_lines)))

    for idx in range(total_lines):
        row = table.rows[start_row + idx]
        if len(row.cells) >= 1:
            row.cells[0].text = op_lines[idx] if idx < len(op_lines) else ""
        if len(row.cells) >= 2:
            row.cells[1].text = std_lines[idx] if idx < len(std_lines) else ""
    return True


def fill_docx_sections(
    docx_input: str,
    docx_output: str,
    question_answers: list[str],
    summary: str | None,
    diagram: str | None,
    placeholder_replacements: dict[str, str] | None = None,
    extra_sections: dict[str, str] | None = None,
) -> None:
    doc = Document(docx_input)
    analysis = analyze_docx(docx_input)

    if placeholder_replacements:

        def replace_text(text: str) -> str:
            for key, value in placeholder_replacements.items():
                text = text.replace(f"{{{{{key}}}}}", value)
            return text

        for p in doc.paragraphs:
            if p.text:
                p.text = replace_text(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        cell.text = replace_text(cell.text)

    for q_idx, slot in enumerate(analysis.questions):
        if q_idx >= len(question_answers):
            break
        table = doc.tables[slot.table_idx]
        answer_cell = table.rows[slot.answer_row_idx].cells[slot.col_idx]
        answer_cell.text = question_answers[q_idx]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                normalized = _normalize(text)

                if summary and normalized.startswith("resumen"):
                    cell.text = f"Resumen\n\n{summary.strip()}"

                if diagram and normalized.startswith("diagrama"):
                    cell.text = f"Diagrama\n\n{diagram.strip()}"

    section_map = {
        "apellidos y nombres": "student_name",
        "id": "student_id",
        "dirección zonal/cfp": "student_address",
        "direccion zonal/cfp": "student_address",
        "carrera": "student_career",
        "curso/ mod. formativo": "student_course",
        "curso/mod. formativo": "student_course",
        "tema de trabajo final": "student_topic",
        "identifica la problemática del caso práctico propuesto": "problem_statement",
        "identifica propuesta de solución y evidencias": "solution_evidence",
        "cronograma de actividades": "schedule",
        "máquinas y equipos": "machines_equipment",
        "maquinas y equipos": "machines_equipment",
        "herramientas e instrumentos": "tools_instruments",
        "materiales e insumos": "materials_supplies",
        "propuesta de solución": "solution_proposal",
        "operaciones / pasos / subpasos": "operations_steps",
        "normas técnicas": "standards_safety_environment",
        "normas tecnicas": "standards_safety_environment",
        "dibujo / esquema / diagrama de propuesta": "textual_diagram",
        "verificar el cumplimiento": "compliance_control",
        "califica el impacto": "evaluation_scores",
    }

    if extra_sections:
        # Intento prioritario: rellenar tablas de "cuadros" por filas y columnas.
        for table in doc.tables:
            if _table_contains_label(table, "cronograma de actividades"):
                _fill_schedule_table(table, str(extra_sections.get("schedule", "")).strip())
            elif _table_contains_label(table, "máquinas y equipos") or _table_contains_label(table, "maquinas y equipos"):
                _fill_resource_table(table, str(extra_sections.get("machines_equipment", "")).strip())
            elif _table_contains_label(table, "herramientas e instrumentos"):
                _fill_resource_table(table, str(extra_sections.get("tools_instruments", "")).strip())
            elif _table_contains_label(table, "materiales e insumos"):
                _fill_resource_table(table, str(extra_sections.get("materials_supplies", "")).strip())
            elif _table_contains_label(table, "operaciones / pasos / subpasos"):
                _fill_execution_table(
                    table,
                    str(extra_sections.get("operations_steps", "")).strip(),
                    str(extra_sections.get("standards_safety_environment", "")).strip(),
                )

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    raw_text = (cell.text or "").strip()
                    if not raw_text:
                        continue
                    normalized = _normalize(raw_text)

                    for label, key in section_map.items():
                        if not _label_matches_cell(label, normalized, key):
                            continue
                        if extra_sections.get(key):
                            value = str(extra_sections[key]).strip()
                            if not value:
                                continue

                            # Campos de estudiante: reemplazar solo la línea correspondiente
                            # para no destruir el formato de celdas con múltiples líneas.
                            if key.startswith("student_"):
                                updated = _replace_student_line(raw_text, value, key)
                                if updated != raw_text:
                                    cell.text = updated
                            else:
                                # Secciones amplias: mantener encabezado y añadir contenido.
                                if value not in raw_text:
                                    cell.text = f"{raw_text}\n\n{value}"
                            break

    output_path = Path(docx_output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
